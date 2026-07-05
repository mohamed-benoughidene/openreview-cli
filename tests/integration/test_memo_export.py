"""Integration tests for memo export CLI command and MemoExporter.

Tests cover:
- T011: Markdown export integration (end-to-end: review runs, .md file created)
- T021: Multiple formats produced in single export
- T021: Duplicate format flag deduplication
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

from docx import Document as DocxDocument

from openreview_cli.review.memo.exporter import MemoExporter
from openreview_cli.review.memo.models import MemoFormat
from openreview_cli.review.models import (
    ClauseAssessment,
    DocMeta,
    Position,
    QAVerdict,
    ReviewReport,
    ReviewSummary,
)


def _make_assessment(
    cid: str,
    pos: Position = Position.PREFERRED,
    conf: float = 0.9,
    color_str: str = "green",
) -> ClauseAssessment:
    ca = ClauseAssessment(
        clause_id=cid,
        clause_text=f"Clause {cid} text for integration testing.",
        playbook_category="confidentiality-term",
        position=pos,
        confidence=conf,
        citation=f"clause {cid} citation",
        qa_verdict=QAVerdict.agree,
        extraction_model="m1",
        qa_model="m1",
    )
    from openreview_cli.review.colors import AssessmentColor

    ca.color = AssessmentColor(color_str)
    return ca


def _make_report(assessments: list[ClauseAssessment] | None = None) -> ReviewReport:
    if assessments is None:
        assessments = [
            _make_assessment("c1", Position.PREFERRED, 0.92, "green"),
            _make_assessment("c2", Position.ACCEPTABLE, 0.78, "amber"),
            _make_assessment("c3", Position.WALKAWAY, 0.55, "red"),
        ]
    dm = DocMeta(
        filename="nda-sample.pdf",
        page_count=5,
        clause_count=len(assessments),
        pii_stripped=True,
    )
    summary = ReviewSummary(
        preferred_count=1,
        acceptable_count=1,
        walkaway_count=1,
        uncertain_count=0,
        no_match_count=0,
        amber_count=1,
        green_count=1,
        red_count=1,
        avg_confidence=sum(a.confidence for a in assessments) / max(len(assessments), 1),
    )
    from datetime import UTC, datetime

    return ReviewReport(
        document=dm,
        assessments=assessments,
        summary=summary,
        playbook_id="precheck-nda-v1",
        generated_at=datetime.now(UTC),
        playbook_version=1,
    )


class TestMarkdownExportIntegration:
    """End-to-end markdown export (T011)."""

    def test_markdown_file_created(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_dir = Path(tmp)
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=output_dir,
                formats={MemoFormat.MARKDOWN},
            )
            result = exporter.export()
            assert MemoFormat.MARKDOWN in result
            path = result[MemoFormat.MARKDOWN]
            assert path.exists()
            assert path.suffix == ".md"

    def test_markdown_contains_required_sections(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.MARKDOWN},
            )
            result = exporter.export()
            content = result[MemoFormat.MARKDOWN].read_text()
            assert "# Memo Export: precheck" in content
            assert "nda-sample.pdf" in content
            assert "precheck-nda-v1" in content
            assert "## Summary" in content
            assert "## Clause Assessments" in content
            assert "## Recommendation" in content
            assert "## Differences" in content
            assert "Disclaimer" in content

    def test_markdown_contains_badges(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.MARKDOWN},
            )
            result = exporter.export()
            content = result[MemoFormat.MARKDOWN].read_text()
            assert "✅" in content
            assert "⚠️" in content
            assert "❌" in content

    def test_markdown_contains_confidence_bars(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.MARKDOWN},
            )
            result = exporter.export()
            content = result[MemoFormat.MARKDOWN].read_text()
            assert "█" in content
            assert "░" in content


class TestJsonExportIntegration:
    """End-to-end JSON export."""

    def test_json_file_created(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.JSON},
            )
            result = exporter.export()
            assert MemoFormat.JSON in result
            path = result[MemoFormat.JSON]
            assert path.exists()
            assert path.suffix == ".json"

    def test_json_valid_and_has_keys(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.JSON},
            )
            result = exporter.export()
            data = json.loads(result[MemoFormat.JSON].read_text())
            assert data["memo_version"] == "1.0"
            assert "clauses" in data
            assert len(data["clauses"]) > 0
            for c in data["clauses"]:
                assert all(k in c for k in ["id", "assessment", "color", "confidence", "citation"])


class TestDocxExportIntegration:
    """End-to-end DOCX export."""

    def test_docx_file_created(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.DOCX},
            )
            result = exporter.export()
            assert MemoFormat.DOCX in result
            path = result[MemoFormat.DOCX]
            assert path.exists()
            assert path.suffix == ".docx"

    def test_docx_has_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.DOCX},
            )
            result = exporter.export()
            doc = DocxDocument(str(result[MemoFormat.DOCX]))
            assert len(doc.tables) > 0

    def test_docx_has_disclaimer(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.DOCX},
            )
            result = exporter.export()
            doc = DocxDocument(str(result[MemoFormat.DOCX]))
            disclaimer_found = any("Disclaimer" in p.text for p in doc.paragraphs)
            assert disclaimer_found


class TestMultiFormatExport:
    """Multiple formats in single export (T021)."""

    def test_all_formats_produced(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.MARKDOWN, MemoFormat.JSON, MemoFormat.DOCX},
            )
            result = exporter.export()
            assert len(result) == 3
            for fmt in (MemoFormat.MARKDOWN, MemoFormat.JSON, MemoFormat.DOCX):
                assert fmt in result
                assert result[fmt].exists()

    def test_duplicate_format_dedup(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            report = _make_report()
            exporter = MemoExporter(
                report=report,
                mode="precheck",
                output_dir=Path(tmp),
                formats={MemoFormat.MARKDOWN, MemoFormat.MARKDOWN},
            )
            result = exporter.export()
            assert len(result) == 1
            assert MemoFormat.MARKDOWN in result


class TestEmptyReport:
    """Edge case: empty report raises error."""

    def test_raises_on_empty_assessments(self) -> None:
        report = _make_report(assessments=[])
        exporter = MemoExporter(report=report, mode="precheck")
        import pytest

        with pytest.raises(ValueError, match="No review results"):
            exporter.export()
