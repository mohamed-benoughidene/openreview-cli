from pathlib import Path
from typing import Any

import pytest

FIXTURES = Path(__file__).resolve().parent.parent / "fixtures" / "docx"


class TestGetHeadingLevel:
    def test_heading_1_returns_0(self) -> None:
        from openreview_cli.parsing.docx_parser import get_heading_level

        class MockStyle:
            name = "Heading 1"

        class MockPara:
            style = MockStyle()
            runs: list[Any] = []

        assert get_heading_level(MockPara()) == 0

    def test_heading_2_returns_1(self) -> None:
        from openreview_cli.parsing.docx_parser import get_heading_level

        class MockStyle:
            name = "Heading 2"

        class MockPara:
            style = MockStyle()
            runs: list[Any] = []

        assert get_heading_level(MockPara()) == 1

    def test_normal_returns_none(self) -> None:
        from openreview_cli.parsing.docx_parser import get_heading_level

        class MockStyle:
            name = "Normal"

        class MockPara:
            style = MockStyle()
            runs: list[Any] = []

        assert get_heading_level(MockPara()) is None


class TestDetectTrackedChanges:
    def test_detects_ins_and_del(self) -> None:
        from openreview_cli.parsing.docx_parser import detect_tracked_changes
        from docx import Document
        result = detect_tracked_changes(Document(str(FIXTURES / "tracked_changes.docx")))
        assert result is True

    def test_no_tracked_changes_returns_false(self) -> None:
        from openreview_cli.parsing.docx_parser import detect_tracked_changes
        from docx import Document
        result = detect_tracked_changes(Document(str(FIXTURES / "simple_contract.docx")))
        assert result is False


class TestSkipEmbeddedImages:
    def test_skips_image_paragraphs(self) -> None:
        from openreview_cli.parsing.docx_parser import skip_embedded_images
        from docx import Document
        doc = Document(str(FIXTURES / "with_images.docx"))
        paras = list(skip_embedded_images(doc.paragraphs))
        texts = [p.text for p in paras]
        assert any("defines key terms" in t for t in texts)
        assert any("follows the embedded image" in t for t in texts)


class TestSourceParagraphIndexing:
    def test_paragraph_index_preserved(self) -> None:
        from openreview_cli.parsing.docx_parser import DocxParser
        clauses = list(DocxParser(FIXTURES / "simple_contract.docx").parse())
        for clause in clauses:
            assert clause.source_paragraph is not None
            assert clause.source_paragraph >= 0
