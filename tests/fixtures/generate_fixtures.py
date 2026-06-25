"""Generate synthetic test fixtures for document parsing phase."""

import os
import shutil
from pathlib import Path
from typing import Any

HERE = Path(__file__).parent

SIMPLE_CONTRACT_TEXT = """Article I: Definitions

Section 1.1: Confidential Information

"Confidential Information" means any information disclosed by the Disclosing Party to the Receiving Party that is marked confidential or that a reasonable person would understand to be confidential given the nature of the information and circumstances of disclosure.

Section 1.2: Receiving Party

The Receiving Party shall hold the Confidential Information in confidence and shall not disclose such information to any third party without prior written consent from the Disclosing Party.

Section 1.3: Disclosing Party

The Disclosing Party retains all right, title, and interest in and to the Confidential Information.

Article II: Obligations

Section 2.1: Non-disclosure

The Receiving Party shall not disclose Confidential Information to any third party. This obligation survives termination of this Agreement.

Section 2.2: Standard of Care

The Receiving Party shall protect Confidential Information using the same degree of care used to protect its own confidential information, but no less than reasonable care.

Article III: Term and Termination

Section 3.1: Term

This Agreement shall commence on the Effective Date and continue for a period of three years.

Section 3.2: Termination

Either party may terminate this Agreement upon thirty days written notice to the other party.

Section 3.3: Survival

The obligations of confidentiality shall survive termination of this Agreement for a period of five years.
"""


def generate_pdf_fixtures() -> None:
    """Generate all PDF fixtures using pymupdf."""
    import pymupdf

    pdf_dir = HERE / "pdf"

    def _make_pdf(text: str, filename: str) -> None:
        doc: Any = pymupdf.open()  # type: ignore[no-untyped-call]
        lines = text.split("\n")
        page = doc.new_page()
        y = 50
        for line in lines:
            if y > 780:
                page = doc.new_page()
                y = 50
            page.insert_text((50, y), line, fontname="helv", fontsize=10)
            y += 15
        doc.save(str(pdf_dir / filename))
        doc.close()

    def _make_multi_column() -> None:
        doc: Any = pymupdf.open()  # type: ignore[no-untyped-call]
        page = doc.new_page()
        cols = [
            "Left Column Paragraph 1: This is the first paragraph in the left column.",
            "Left Column Paragraph 2: This is the second paragraph in the left column.",
            "Left Column Paragraph 3: More content in the left column area.",
            "Right Column Paragraph 1: This text appears in the right column area.",
            "Right Column Paragraph 2: Second paragraph in the right column area.",
            "Right Column Paragraph 3: More content positioned in the right column.",
        ]
        for i, text in enumerate(cols):
            x = 50 if i < 3 else 300
            y = 50 + (i % 3) * 200
            page.insert_text((x, y), text, fontname="helv", fontsize=10)
        doc.save(str(pdf_dir / "multi_column.pdf"))
        doc.close()

    _make_pdf(SIMPLE_CONTRACT_TEXT, "simple_contract.pdf")

    complex_text = """Article I: Definitions

Section 1.1: Confidential Information

"As used in this Agreement, Confidential Information means all information disclosed by one party to the other."

Section 1.2: Exclusions

(a) Information that is or becomes publicly available without breach of this Agreement.
(b) Information that the Receiving Party can demonstrate was known prior to disclosure.
(i) Any combination of the above items is also excluded.
(ii) This exclusion does not apply to trade secrets.

Article II: Scope

Section 2.1: Coverage

This Agreement covers all written and oral communications between the parties.

Section 2.2: Exceptions

(1) Information obtained from a third party without restriction.
(2) Information independently developed by the Receiving Party.
(3) Information required to be disclosed by law.

Article III: Miscellaneous

Section 3.1: Governing Law

This Agreement shall be governed by the laws of the State of Delaware.

Section 3.2: Dispute Resolution

Any dispute arising under this Agreement shall be resolved through binding arbitration.
"""
    _make_pdf(complex_text, "complex_numbering.pdf")

    flat_text = """The parties agree to the following terms and conditions.

This Agreement shall be governed by the laws of the State of New York.

Any disputes arising under this Agreement shall be resolved through binding arbitration.

Each party shall indemnify the other party against any losses arising from breach of this Agreement.

This Agreement represents the entire understanding between the parties.

No modification of this Agreement shall be effective unless in writing and signed by both parties.
"""
    _make_pdf(flat_text, "flat_document.pdf")

    _make_multi_column()

    plain_doc: Any = pymupdf.open()  # type: ignore[no-untyped-call]
    plain_page = plain_doc.new_page()
    plain_page.insert_text((50, 50), "Hello World", fontname="helv", fontsize=12)
    plain_doc.save(str(pdf_dir / "password_protected.pdf"),
                   encryption=pymupdf.PDF_ENCRYPT_AES_128,  # type: ignore[attr-defined]
                   user_pw="test123",
                   owner_pw="test123")
    plain_doc.close()

    good_pdf: Any = pymupdf.open()  # type: ignore[no-untyped-call]
    good_pdf.new_page()
    good_pdf.new_page()
    good_pdf.new_page()
    good_pdf.save(str(pdf_dir / "_valid_temp.pdf"))
    good_pdf.close()
    data = open(pdf_dir / "_valid_temp.pdf", "rb").read()[:-512]
    open(pdf_dir / "corrupt.pdf", "wb").write(data)
    os.remove(pdf_dir / "_valid_temp.pdf")

    open(pdf_dir / "empty.pdf", "wb").close()

    def _make_large_pdf(page_count: int, filename: str) -> None:
        doc: Any = pymupdf.open()  # type: ignore[no-untyped-call]
        for i in range(page_count):
            page = doc.new_page()
            page.insert_text((50, 50), f"Clause {i+1}: This is a sample clause for testing purposes.", fontname="helv", fontsize=10)
        doc.save(str(pdf_dir / filename))
        doc.close()

    _make_large_pdf(50, "50_page.pdf")
    _make_large_pdf(500, "500_page.pdf")

    print("PDF fixtures generated.")


def generate_docx_fixtures() -> None:
    """Generate all DOCX fixtures using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree  # type: ignore[import-untyped]

    docx_dir = HERE / "docx"
    WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _make_simple_contract() -> None:
        doc = Document()
        for line in SIMPLE_CONTRACT_TEXT.strip().split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("Article") or line.startswith("Section"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            else:
                doc.add_paragraph(line)
        doc.save(str(docx_dir / "simple_contract.docx"))

    def _make_with_headings() -> None:
        doc = Document()
        doc.add_heading("Article I: Definitions", level=1)
        doc.add_paragraph("This Agreement sets forth the definitions applicable hereto.")
        doc.add_heading("Section 1.1: Confidential Information", level=2)
        doc.add_paragraph("Confidential Information means all proprietary information disclosed.")
        doc.add_heading("Section 1.2: Exclusions", level=2)
        doc.add_paragraph("Confidential Information does not include publicly available information.")
        doc.save(str(docx_dir / "with_headings.docx"))

    def _make_tracked_changes() -> None:
        doc = Document()
        doc.add_paragraph("This is the original text.")
        body = doc.element.body
        ins = etree.SubElement(body, f"{{{WORD_NS}}}ins")
        ins.set(f"{{{WORD_NS}}}id", "0")
        ins.set(f"{{{WORD_NS}}}author", "test")
        ins.set(f"{{{WORD_NS}}}date", "2026-01-01T00:00:00Z")
        run = etree.SubElement(ins, f"{{{WORD_NS}}}r")
        t = etree.SubElement(run, f"{{{WORD_NS}}}t")
        t.text = " Inserted text."
        delel = etree.SubElement(body, f"{{{WORD_NS}}}del")
        delel.set(f"{{{WORD_NS}}}id", "1")
        r2 = etree.SubElement(delel, f"{{{WORD_NS}}}r")
        del_text = etree.SubElement(r2, f"{{{WORD_NS}}}delText")
        del_text.text = "Deleted text."
        doc.save(str(docx_dir / "tracked_changes.docx"))

    def _make_flat_document() -> None:
        doc = Document()
        doc.add_paragraph("The parties agree to the following terms and conditions.")
        doc.add_paragraph("This Agreement shall be governed by the laws of the State of New York.")
        doc.add_paragraph("Any disputes arising under this Agreement shall be resolved through binding arbitration.")
        doc.add_paragraph("Each party shall indemnify the other party against any losses.")
        doc.add_paragraph("This Agreement represents the entire understanding between the parties.")
        doc.save(str(docx_dir / "flat_document.docx"))

    def _make_with_images() -> None:
        doc = Document()
        doc.add_heading("Article I: Definitions", level=1)
        doc.add_paragraph("This paragraph defines key terms used in this agreement.")
        p = doc.add_paragraph()
        run = p.add_run(" ")
        drawing = etree.SubElement(p._p, f"{{{WORD_NS}}}drawing")
        doc.add_paragraph("This paragraph follows the embedded image.")
        doc.add_heading("Section 1.1: Key Terms", level=2)
        doc.save(str(docx_dir / "with_images.docx"))

    _make_simple_contract()
    _make_with_headings()
    _make_tracked_changes()
    _make_flat_document()
    _make_with_images()
    print("DOCX fixtures generated.")


if __name__ == "__main__":
    generate_pdf_fixtures()
    generate_docx_fixtures()
